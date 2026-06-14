#!/usr/bin/env python3
"""Render a DOCX file from a JSON payload.

Reads JSON from stdin, writes binary .docx to stdout, errors to stderr.

Payload shape:
{
  "title": "Music Title",
  "author": "Author Name",
  "sections": [
    {
      "type": "verse",
      "lines": [
        {"chord_line": "G  C", "lyric_line": "Hello world"},
        {"chord_line": "", "lyric_line": "Lyric line without chords"}
      ]
    },
    {
      "type": "label",
      "label": "[CORO]"
    }
  ]
}
"""
import json
import sys
import io
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def render(payload):
    songs = payload.get("songs")
    if songs is None:
        # Single song mode
        songs = [payload]
    
    doc = Document()
    
    # Set Narrow Margins (0.1 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.1)
        section.left_margin = Inches(0.1)
        section.right_margin = Inches(0.1)
    
    for i, song in enumerate(songs):
        if i > 0:
            doc.add_page_break()
            
        title = song.get("title", "Untitled")
        sections = song.get("sections", [])
        
        # Title - Bold, Left Aligned
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Roboto Mono"
        run.font.size = Pt(14)
        
        doc.add_paragraph() # Spacer
        
        for section in sections:
            if section.get("type") == "label":
                label = section.get("label", "").strip()
                if label:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(label)
                    run.bold = True
                    run.font.name = "Roboto Mono"
                    run.font.size = Pt(10)
                continue

            for line in section.get("lines", []):
                chords = line.get("chord_line", "").rstrip()
                lyrics = line.get("lyric_line", "").rstrip()
                
                if chords:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(chords)
                    run.bold = True
                    run.font.name = "Roboto Mono"
                    run.font.size = Pt(12)
                
                if lyrics:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(lyrics)
                    run.font.name = "Roboto Mono"
                    run.font.size = Pt(12)
                elif not chords:
                    # Blank line
                    doc.add_paragraph()
            
            doc.add_paragraph() # Space between sections

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def main():
    try:
        payload = json.load(sys.stdin)
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"Invalid JSON payload: {exc}\n")
        return 1

    try:
        binary = render(payload)
    except Exception as exc:
        sys.stderr.write(f"Render failed: {exc}\n")
        import traceback
        traceback.print_exc(file=sys.stderr)
        return 1

    sys.stdout.buffer.write(binary)
    return 0

if __name__ == "__main__":
    sys.exit(main())
